"""
POST /api/resume/extract  — extract plain text from PDF or DOCX upload.

PDF: uses pymupdf (fitz) — handles multi-column, tables, sidebars correctly.
DOCX: uses python-docx.
Both run a cleanup pass to remove extraction artefacts.
"""

import io
import re
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.db.database import get_db
from app.db.crud import ResumeRepository
from app.db.models import User
from app.deps import get_current_user

log = logging.getLogger(__name__)
router = APIRouter(prefix="/resume", tags=["resume"])

MAX_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB


class ExtractResponse(BaseModel):
    text:     str
    chars:    int
    filename: str


class ResumeOut(BaseModel):
    text:     str
    filename: str | None = None


@router.post("/extract", response_model=ExtractResponse)
async def extract_resume(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext == "doc":
        # python-docx only reads the modern XML .docx, not the legacy binary .doc.
        raise HTTPException(400, "Legacy .doc isn't supported — re-save as .docx or PDF.")
    if ext not in ("pdf", "docx"):
        raise HTTPException(400, f"Unsupported format: .{ext}. Use PDF or DOCX.")

    raw = await file.read()
    if len(raw) > MAX_SIZE_BYTES:
        raise HTTPException(413, "File too large. Max 15 MB.")

    try:
        text = _extract_pdf(raw) if ext == "pdf" else _extract_docx(raw)
        text = _clean(text)
    except Exception as exc:
        # pymupdf/python-docx exceptions are library-internal detail (parser
        # state, C-extension messages) — log it, tell the user something they
        # can act on instead.
        log.warning(f"Extraction failed for {file.filename}: {exc}")
        raise HTTPException(422,
            "Couldn't read that file. It may be corrupted, password-protected, "
            "or a scanned image without selectable text — try re-saving it or "
            "pasting the text directly.")

    if not text.strip():
        raise HTTPException(422, "No readable text found. Try saving as a different format.")

    text = text.strip()
    # Persist so compose can default to it without re-uploading.
    ResumeRepository(db, user.id).save(text, file.filename)

    return ExtractResponse(text=text, chars=len(text), filename=file.filename)


class ResumeSaveRequest(BaseModel):
    text: str


@router.post("/save", response_model=ResumeOut)
def save_resume_text(
    req: ResumeSaveRequest,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Save manually pasted resume text (no file upload required)."""
    text = req.text.strip()
    if not text:
        raise HTTPException(400, "Resume text cannot be empty")
    if len(text) > 50_000:
        raise HTTPException(400, "Resume text too long (max 50,000 chars)")
    saved = ResumeRepository(db, user.id).save(text, filename=None)
    return ResumeOut(text=saved.text, filename=saved.filename)


@router.get("/latest", response_model=ResumeOut)
def latest_resume(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """Return the most recently saved résumé for the current user (empty if none)."""
    resume = ResumeRepository(db, user.id).get_latest()
    if not resume:
        return ResumeOut(text="", filename=None)
    return ResumeOut(text=resume.text, filename=resume.filename)


# ── PDF extraction via pymupdf ────────────────────────────────────────────────
def _extract_pdf(data: bytes) -> str:
    import fitz  # pymupdf

    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []

    for page in doc:
        # extract_text with "blocks" sort reads in natural reading order
        # even for multi-column layouts
        blocks = page.get_text("blocks", sort=True)

        page_lines: list[str] = []
        prev_y = None

        for block in blocks:
            # block = (x0, y0, x1, y1, text, block_no, block_type)
            if block[6] != 0:   # skip image blocks
                continue
            block_text = block[4].strip()
            if not block_text:
                continue

            # Add blank line between blocks that are far apart vertically
            if prev_y is not None and (block[1] - prev_y) > 20:
                page_lines.append("")

            page_lines.append(block_text)
            prev_y = block[3]

        pages.append("\n".join(page_lines))

    doc.close()
    return "\n\n".join(pages)


# ── DOCX extraction via python-docx ──────────────────────────────────────────
def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
        else:
            # Preserve intentional blank lines (section breaks)
            if lines and lines[-1] != "":
                lines.append("")

    # Extract table cells in row order
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  |  ".join(cells))

    return "\n".join(lines)


# ── Post-extraction cleanup ───────────────────────────────────────────────────
def _clean(text: str) -> str:
    # Fix spaced-out characters: "E x p e r i e n c e" → "Experience"
    # A run of 3+ single-letter "words" separated by single spaces is virtually
    # always a font-kerning artifact from PDF extraction, never real text.
    text = re.sub(r'\b(?:[A-Za-z] ){2,}[A-Za-z]\b', lambda m: m.group(0).replace(" ", ""), text)

    # Remove ligature / encoding artefacts
    text = text.replace('ﬁ', 'fi').replace('ﬂ', 'fl')
    text = text.replace('’', "'").replace('‘', "'")
    text = text.replace('–', '-').replace('—', '--')
    text = text.replace(' ', ' ')

    # Collapse 3+ consecutive blank lines → 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove lines that are just dots/dashes (PDF table-of-contents artifacts)
    text = re.sub(r'^\s*[.\-_]{3,}\s*$', '', text, flags=re.MULTILINE)

    # Trim trailing whitespace on each line
    text = '\n'.join(line.rstrip() for line in text.splitlines())

    return text.strip()
